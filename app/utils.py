import os
from pathlib import Path
from typing import List

# Lazy-loaded heavy dependencies
_embeddings = None
_Document = None
_RecursiveCharacterTextSplitter = None
_FAISS = None

VECTORSTORE_PATH = Path("vectorstore")


def _get_document_class():
    global _Document
    if _Document is None:
        from langchain_core.documents import Document
        _Document = Document
    return _Document


def _get_text_splitter():
    global _RecursiveCharacterTextSplitter
    if _RecursiveCharacterTextSplitter is None:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        _RecursiveCharacterTextSplitter = RecursiveCharacterTextSplitter
    return _RecursiveCharacterTextSplitter


def _get_faiss():
    global _FAISS
    if _FAISS is None:
        from langchain_community.vectorstores import FAISS
        _FAISS = FAISS
    return _FAISS


def get_embeddings():
    global _embeddings
    if _embeddings is None:
        print("Loading embedding model...")
        from langchain_community.embeddings import HuggingFaceEmbeddings
        _embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return _embeddings


def _extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF using pypdf (lightweight, no system deps)"""
    from pypdf import PdfReader
    try:
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error reading PDF {file_path.name}: {e}")
        return ""


def _extract_text_from_file(file_path: Path) -> str:
    """Extract text from various file types"""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_text_from_pdf(file_path)

    elif suffix in {".txt", ".md", ".html", ".htm"}:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            # Basic HTML tag stripping for HTML files
            if suffix in {".html", ".htm"}:
                import re
                content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
                content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
                content = re.sub(r'<[^>]+>', ' ', content)
                content = re.sub(r'\s+', ' ', content)
            return content.strip()
        except Exception as e:
            print(f"Error reading {file_path.name}: {e}")
            return ""

    elif suffix == ".docx":
        try:
            # Try to use python-docx if available
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except ImportError:
            print(f"python-docx not installed, skipping {file_path.name}")
            return ""
        except Exception as e:
            print(f"Error reading DOCX {file_path.name}: {e}")
            return ""

    return ""


def load_all_documents(folder: str = "data") -> List:
    """Loads PDF, DOCX, TXT, HTML files from data folder using lightweight parsers"""
    Document = _get_document_class()

    docs = []
    folder_path = Path(folder)
    if not folder_path.exists():
        print("No 'data' folder found, create it and drop your files there")
        return docs

    supported_extensions = {".pdf", ".docx", ".txt", ".html", ".htm", ".md"}

    for file_path in folder_path.rglob("*.*"):
        if file_path.suffix.lower() in supported_extensions:
            print(f"Loading ---> {file_path.name}")
            text = _extract_text_from_file(file_path)
            if text.strip():
                docs.append(
                    Document(
                        page_content=text.strip(),
                        metadata={"source": file_path.name}
                    )
                )

    print(f"Loaded {len(docs)} documents")
    return docs


def chunk_documents(docs: List) -> List:
    RecursiveCharacterTextSplitter = _get_text_splitter()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = splitter.split_documents(docs)
    print(f"Split into {len(chunks)} chunks")
    return chunks


def create_and_store_vectorstore() -> None:
    FAISS = _get_faiss()
    raw_docs = load_all_documents()
    if not raw_docs:
        return
    chunks = chunk_documents(raw_docs)
    vectorstore = FAISS.from_documents(chunks, get_embeddings())
    VECTORSTORE_PATH.mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(str(VECTORSTORE_PATH))
    print(f"Vectorstore created and saved in '{VECTORSTORE_PATH}'")


def get_retriever():
    """Going to be used by AGENT, loads existing vectorstore or creates it once"""
    FAISS = _get_faiss()

    if VECTORSTORE_PATH.exists():
        print("Loading existing vectorstore")
        vectorstore = FAISS.load_local(
            str(VECTORSTORE_PATH),
            get_embeddings(),
            allow_dangerous_deserialization=True
        )
    else:
        print("First run detected, building vectorstore")
        create_and_store_vectorstore()
        vectorstore = FAISS.load_local(str(VECTORSTORE_PATH), get_embeddings(), allow_dangerous_deserialization=True)

    return vectorstore.as_retriever(search_kwargs={"k": 10})


__all__ = ["get_retriever"]
